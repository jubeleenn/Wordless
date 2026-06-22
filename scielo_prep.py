# -*- coding: utf-8 -*-
import tkinter as tk
from tkinter import filedialog, messagebox
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

class ScieloPrepAppFull:
    def __init__(self, root):
        self.root = root
        self.root.title("SciELO Prep PRO | Automatización y Reportes")
        self.root.geometry("500x350")
        self.root.configure(bg="#fdfdfc")
        self.word_path = None

        tk.Label(root, text="Preparación SciELO PRO", font=("Arial", 16, "bold"), bg="#fdfdfc", fg="#012662").pack(pady=(20, 5))
        tk.Label(root, text="Aplica TODAS las reglas y genera un reporte de cambios", font=("Arial", 10), bg="#fdfdfc", fg="#5a6a7a").pack(pady=5)

        self.lbl_archivo = tk.Label(root, text="Ningún archivo seleccionado", fg="red", bg="#fdfdfc")
        self.lbl_archivo.pack(pady=10)

        tk.Button(root, text="📂 Seleccionar Word (.docx)", command=self.seleccionar_archivo, font=("Arial", 10)).pack(pady=5)
        
        self.btn_procesar = tk.Button(root, text="⚙️ Aplicar Marcación y Generar Reporte", font=("Arial", 12, "bold"), bg="#1e6292", fg="white", command=self.procesar_documento)
        self.btn_procesar.pack(pady=20, fill="x", padx=40)

    def seleccionar_archivo(self):
        ruta = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if ruta:
            self.word_path = ruta
            self.lbl_archivo.config(text=Path(ruta).name, fg="green")

    def quitar_hipervinculos(self, doc, log_cambios):
        """Elimina hipervínculos a nivel XML conservando el texto y documentando el cambio"""
        for i, p in enumerate(doc.paragraphs):
            hyperlinks = p._element.findall('.//w:hyperlink', doc.part.element.nsmap)
            if hyperlinks:
                log_cambios.append(f"Párrafo {i+1}: Se detectaron y eliminaron {len(hyperlinks)} hipervínculo(s) conservando el texto intacto (Regla SciELO).")
                for hl in hyperlinks:
                    index = p._element.index(hl)
                    runs = hl.findall('.//w:r', doc.part.element.nsmap)
                    for r in reversed(runs):
                        p._element.insert(index, r)
                    p._element.remove(hl)

    def procesar_documento(self):
        if not self.word_path:
            messagebox.showwarning("Atención", "Por favor seleccioná un archivo Word.")
            return

        try:
            doc = docx.Document(self.word_path)
            
            # --- INICIALIZAMOS EL REPORTE DE CAMBIOS ---
            log_cambios = []
            log_cambios.append(f"--- REPORTE DE PREPARACIÓN SCIELO ---")
            log_cambios.append(f"Archivo original procesado: {Path(self.word_path).name}\n")
            log_cambios.append("REGLA GENERAL APLICADA: Todo el cuerpo del texto fue llevado a tamaño 12pts y alineación Justificada (excepto excepciones listadas abajo).\n")
            log_cambios.append("DETALLE DE CAMBIOS POR PÁRRAFO ORIGINAL:\n" + "-"*50)
            
            # 1. Eliminar Hipervínculos de forma segura
            self.quitar_hipervinculos(doc, log_cambios)
            
            parrafos_a_borrar = []

            # 2. Recorrer y formatear
            for i, p in enumerate(doc.paragraphs):
                texto_limpio = p.text.strip()
                texto_lower = texto_limpio.lower()

                # Eliminar párrafos vacíos (Regla SciELO: sin saltos de línea)
                if not texto_limpio:
                    parrafos_a_borrar.append(p)
                    log_cambios.append(f"Párrafo {i+1}: ELIMINADO (Párrafo en blanco / salto de línea innecesario).")
                    continue

                # Setear tamaño 12 a todo como base
                for run in p.runs:
                    run.font.size = Pt(12)

                # Regla DOI o Sección
                if texto_lower.startswith("doi:") or (i < 3 and len(texto_limpio) < 50 and not texto_lower.startswith("resumen")):
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    log_cambios.append(f"Párrafo {i+1}: Detectado como DOI o Sección. Alineado a la derecha.")

                # Regla Resumen / Abstract
                elif texto_lower.startswith("resumen") or texto_lower.startswith("abstract"):
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if ":" in texto_limpio:
                        partes = texto_limpio.split(":", 1)
                        p.clear()
                        # CORRECCIÓN AQUÍ
                        r1 = p.add_run(partes[0] + ":") 
                        r1.bold = True
                        r1.font.size = Pt(12)
                        if len(partes) > 1:
                            r2 = p.add_run(partes[1])
                            r2.font.size = Pt(12)
                    else:
                        p.clear()
                        r = p.add_run(texto_limpio)
                        r.bold = True
                        r.font.size = Pt(12)
                    log_cambios.append(f"Párrafo {i+1}: Detectado como Resumen/Abstract. Alineado a la izquierda y prefijo en Negrita.")

                # Regla Palabras Clave / Key Words
                elif texto_lower.startswith("palabras clave") or texto_lower.startswith("key words"):
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if ":" in texto_limpio:
                        partes = texto_limpio.split(":", 1)
                        p.clear()
                        # CORRECCIÓN AQUÍ
                        r1 = p.add_run(partes[0] + ":") 
                        r1.bold = True
                        r1.font.size = Pt(12)
                        if len(partes) > 1:
                            r2 = p.add_run(partes[1])
                            r2.font.size = Pt(12)
                    log_cambios.append(f"Párrafo {i+1}: Detectado como Palabras Clave. Alineado a la izquierda y prefijo en Negrita.")

                # Regla Agradecimientos
                elif texto_lower in ["agradecimientos", "agradecimiento"]:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.clear()
                    r = p.add_run(texto_limpio)
                    r.bold = True
                    r.font.size = Pt(12)
                    p.insert_paragraph_before("") # Línea en blanco antes
                    log_cambios.append(f"Párrafo {i+1}: Detectado como Agradecimientos. Centrado, Negrita y se agregó línea en blanco previa.")
                    
                # Regla Referencias Bibliográficas
                elif "referencias bibliográficas" in texto_lower or "referencias" == texto_lower:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.clear()
                    r = p.add_run(texto_limpio)
                    r.bold = True
                    r.font.size = Pt(12)
                    log_cambios.append(f"Párrafo {i+1}: Detectado como Título de Referencias. Alineado a la izquierda y Negrita.")

                # Detección de Títulos y Subtítulos (Cortos, negrita original, sin punto final)
                elif len(texto_limpio) < 100 and not texto_limpio.endswith(".") and any(r.bold for r in p.runs):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.insert_paragraph_before("") # Salto antes
                    
                    # Si es todo mayúsculas o es de los primeros, asumimos Título Principal (16pts)
                    if texto_limpio.isupper() or i < 10:
                        for r in p.runs:
                            r.font.size = Pt(16)
                            r.bold = True
                        p.insert_paragraph_before("") # Dos saltos antes (SciELO)
                        log_cambios.append(f"Párrafo {i+1}: Detectado como Título Principal. Centrado, tamaño 16pts y se aislaron líneas en blanco previas.")
                    else:
                        # Subtítulo (14pts)
                        for r in p.runs:
                            r.font.size = Pt(14)
                            r.bold = True
                        log_cambios.append(f"Párrafo {i+1}: Detectado como Subtítulo. Centrado, tamaño 14pts y se aisló con línea en blanco.")

                # Detección de Citas Textuales Largas (> 3 líneas aprox, comillas o ya indentadas)
                elif len(texto_limpio) > 250 and (texto_limpio.startswith('"') or p.paragraph_format.left_indent):
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.left_indent = Cm(4) # Regla SciELO: 4 cm izquierda
                    log_cambios.append(f"Párrafo {i+1}: Detectado como Cita Textual Larga. Se aplicó sangría izquierda estricta de 4 cm y justificado.")

                # Resto del cuerpo del texto
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Borrar los párrafos vacíos que marcamos
            for p in parrafos_a_borrar:
                p._element.getparent().remove(p._element)

            # 3. Guardar el archivo Word procesado
            ruta_salida_docx = Path(self.word_path).with_name(Path(self.word_path).stem + "_PRO_SciELO.docx")
            doc.save(ruta_salida_docx)

            # 4. Guardar el Reporte en TXT
            ruta_salida_txt = Path(self.word_path).with_name(Path(self.word_path).stem + "_REPORTE_Cambios.txt")
            with open(ruta_salida_txt, "w", encoding="utf-8") as f:
                f.write("\n".join(log_cambios))

            messagebox.showinfo("¡Trabajo Completado!", 
                                f"El archivo fue procesado con éxito corrigiendo el formato.\n\n"
                                f"Se han generado DOS archivos nuevos en tu carpeta:\n\n"
                                f"📄 1. El Word listo: {ruta_salida_docx.name}\n"
                                f"📋 2. El Reporte: {ruta_salida_txt.name}\n\n"
                                f"¡Revisá el reporte para ver qué cambió exactamente la app!")

        except Exception as e:
            messagebox.showerror("Error", f"Ocurrió un problema procesando el archivo:\n{str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    app = ScieloPrepAppFull(root)
    root.mainloop()